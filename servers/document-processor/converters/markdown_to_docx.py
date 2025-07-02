"""
Markdown to DOCX converter implementation.
"""

import time
from pathlib import Path
from typing import Optional
import markdown
from datetime import datetime
import re

from domain import (
    Document, DocumentFormat, ConversionOptions, ConversionResult, 
    DocumentConverter
)
from domain.exceptions import ConversionError
from .mermaid_processor import preprocess_mermaid


class MarkdownToDocxConverter(DocumentConverter):
    """
    Converts markdown documents to DOCX format.
    
    This converter creates Microsoft Word documents with proper formatting.
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize the converter.
        
        Args:
            output_dir: Directory for output files
        """
        self.output_dir = output_dir or Path("output")
        
        # Ensure directory exists
        self.output_dir.mkdir(exist_ok=True)
    
    @property
    def source_format(self) -> DocumentFormat:
        """The source format this converter accepts."""
        return DocumentFormat.MARKDOWN
    
    @property
    def target_format(self) -> DocumentFormat:
        """The target format this converter produces."""
        return DocumentFormat.DOCX
    
    @property
    def name(self) -> str:
        """Human-readable name of this converter."""
        return "Markdown to DOCX Converter"
    
    def convert(self, document: Document, options: ConversionOptions) -> ConversionResult:
        """
        Convert a markdown document to DOCX.
        
        Args:
            document: The source markdown document
            options: Conversion options and settings
            
        Returns:
            ConversionResult with the conversion outcome
        """
        start_time = time.time()
        
        try:
            # Validate the document
            self.validate_document(document)
            
            # Step 1: Parse markdown and create DOCX
            output_file_path = self._generate_output_path(document, options)
            self._markdown_to_docx(document, output_file_path, options)
            
            # Calculate conversion time
            conversion_time = time.time() - start_time
            
            return ConversionResult.success_result(
                output_file_path=output_file_path,
                target_format=self.target_format,
                engine_used="python-docx",
                conversion_time_seconds=conversion_time,
                metadata={
                    "original_word_count": document.word_count,
                    "original_char_count": document.character_count,
                }
            )
            
        except Exception as e:
            conversion_time = time.time() - start_time
            
            return ConversionResult.error_result(
                error_message=str(e),
                error_details=f"Conversion failed after {conversion_time:.2f} seconds",
                target_format=self.target_format,
                conversion_time_seconds=conversion_time
            )
    
    def _markdown_to_docx(self, document: Document, output_path: Path, options: ConversionOptions) -> None:
        """Convert markdown content to DOCX."""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            # Step 1: Preprocess Mermaid diagrams
            output_dir = Path(options.output_directory) if options.output_directory else self.output_dir
            processed_content = preprocess_mermaid(
                document.content, 
                output_dir=output_dir,
                image_format="png",  # PNG works well in DOCX
                theme="default"
            )
            
            # Create new document
            doc = DocxDocument()
            
            # Set document properties
            properties = doc.core_properties
            if options.title or document.metadata.title:
                properties.title = options.title or document.metadata.title
            if options.author or document.metadata.author:
                properties.author = options.author or document.metadata.author
            if options.subject or document.metadata.subject:
                properties.subject = options.subject or document.metadata.subject
            
            # Parse processed markdown content line by line
            lines = processed_content.split('\n')
            i = 0
            
            while i < len(lines):
                line = lines[i].rstrip()
                
                if not line:
                    # Empty line - add paragraph break
                    doc.add_paragraph()
                    i += 1
                    continue
                
                # Headers
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('# ').strip()
                    if level == 1:
                        heading = doc.add_heading(text, level=1)
                    elif level == 2:
                        heading = doc.add_heading(text, level=2)
                    elif level == 3:
                        heading = doc.add_heading(text, level=3)
                    else:
                        heading = doc.add_heading(text, level=4)
                
                # Code blocks
                elif line.startswith('```'):
                    i += 1
                    code_lines = []
                    while i < len(lines) and not lines[i].startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    
                    # Add code block
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Inches(0.12)
                
                # Blockquotes
                elif line.startswith('>'):
                    quote_lines = []
                    while i < len(lines) and lines[i].startswith('>'):
                        quote_lines.append(lines[i][1:].strip())
                        i += 1
                    i -= 1  # Back up one since we'll increment at end
                    
                    quote_text = ' '.join(quote_lines)
                    p = doc.add_paragraph(quote_text)
                    p.style = 'Quote'
                
                # Lists
                elif line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', line):
                    list_items = []
                    is_numbered = re.match(r'^\d+\. ', line)
                    
                    while i < len(lines):
                        current_line = lines[i].rstrip()
                        if current_line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\. ', current_line):
                            # Extract list item text
                            if is_numbered:
                                item_text = re.sub(r'^\d+\. ', '', current_line)
                            else:
                                item_text = current_line[2:]  # Remove '- ', '* ', or '+ '
                            list_items.append(item_text)
                            i += 1
                        elif current_line == '':
                            i += 1
                            break
                        else:
                            break
                    i -= 1  # Back up one since we'll increment at end
                    
                    # Add list items
                    for item in list_items:
                        p = doc.add_paragraph(item, style='List Bullet' if not is_numbered else 'List Number')
                
                # Tables (basic support)
                elif '|' in line and line.count('|') >= 2:
                    table_lines = []
                    while i < len(lines) and '|' in lines[i]:
                        table_lines.append(lines[i])
                        i += 1
                    i -= 1  # Back up one
                    
                    if len(table_lines) >= 2:
                        # Parse table
                        rows = []
                        for table_line in table_lines:
                            if '---' in table_line:  # Skip separator line
                                continue
                            cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                            if cells:
                                rows.append(cells)
                        
                        if rows:
                            # Create table
                            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                            table.style = 'Table Grid'
                            
                            for row_idx, row_data in enumerate(rows):
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < len(table.rows[row_idx].cells):
                                        table.rows[row_idx].cells[col_idx].text = cell_data
                
                # Image references (![alt text](path))
                elif line.startswith('![') and '](' in line and line.endswith(')'):
                    # Extract image path
                    alt_start = line.find('[') + 1
                    alt_end = line.find(']')
                    path_start = line.find('](') + 2
                    path_end = line.rfind(')')
                    
                    if alt_start < alt_end and path_start < path_end:
                        alt_text = line[alt_start:alt_end]
                        image_path = line[path_start:path_end]
                        
                        # Handle file:// URLs
                        if image_path.startswith('file://'):
                            image_path = image_path[7:]  # Remove file:// prefix
                        
                        # Add image to document
                        try:
                            from pathlib import Path as PathLib
                            img_path = PathLib(image_path)
                            if img_path.exists():
                                p = doc.add_paragraph()
                                run = p.add_run()
                                run.add_picture(str(img_path), width=Inches(6))  # 6 inch width
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                # Fallback: add alt text if image not found
                                p = doc.add_paragraph(f"[Image: {alt_text}]")
                        except Exception as e:
                            # Fallback: add alt text if image embedding fails
                            p = doc.add_paragraph(f"[Image: {alt_text}]")
                
                # Regular paragraphs
                else:
                    # Handle inline formatting
                    text = self._process_inline_formatting(line)
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, text)
                
                i += 1
            
            # Save document
            doc.save(str(output_path))
            
        except ImportError:
            raise ConversionError(
                source_format=self.source_format.value,
                target_format=self.target_format.value,
                message="python-docx not available",
                details="Install python-docx: pip install python-docx"
            )
        except Exception as e:
            raise ConversionError(
                source_format=self.source_format.value,
                target_format=self.target_format.value,
                message="DOCX conversion failed",
                details=str(e)
            )
    
    def _process_inline_formatting(self, text: str) -> list:
        """Process inline markdown formatting and return list of (text, format) tuples."""
        # This is a simplified implementation
        # In a full implementation, you'd want a proper markdown parser
        
        parts = []
        current_text = text
        
        # Handle bold (**text**)
        while '**' in current_text:
            start = current_text.find('**')
            if start == -1:
                break
            end = current_text.find('**', start + 2)
            if end == -1:
                break
            
            # Add text before bold
            if start > 0:
                parts.append((current_text[:start], 'normal'))
            
            # Add bold text
            bold_text = current_text[start + 2:end]
            parts.append((bold_text, 'bold'))
            
            # Continue with remaining text
            current_text = current_text[end + 2:]
        
        # Handle italic (*text*)
        while '*' in current_text and '**' not in current_text:
            start = current_text.find('*')
            if start == -1:
                break
            end = current_text.find('*', start + 1)
            if end == -1:
                break
            
            # Add text before italic
            if start > 0:
                parts.append((current_text[:start], 'normal'))
            
            # Add italic text
            italic_text = current_text[start + 1:end]
            parts.append((italic_text, 'italic'))
            
            # Continue with remaining text
            current_text = current_text[end + 1:]
        
        # Add remaining text
        if current_text:
            parts.append((current_text, 'normal'))
        
        return parts if parts else [(text, 'normal')]
    
    def _add_formatted_text(self, paragraph, text_parts):
        """Add formatted text to a paragraph."""
        if isinstance(text_parts, str):
            text_parts = [(text_parts, 'normal')]
        
        for text, format_type in text_parts:
            run = paragraph.add_run(text)
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'italic':
                run.italic = True
    
    def _generate_output_path(self, document: Document, options: ConversionOptions) -> Path:
        """Generate the output file path."""
        if options.filename:
            filename = options.filename
        elif document.metadata.title:
            # Use document title, sanitized for filename
            filename = "".join(c for c in document.metadata.title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = filename.replace(' ', '_')
        else:
            # Generate timestamp-based filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}"
        
        # Ensure .docx extension
        if not filename.endswith('.docx'):
            filename += '.docx'
        
        # Use custom output directory if specified
        output_dir = Path(options.output_directory) if options.output_directory else self.output_dir
        output_dir.mkdir(parents=True, exist_ok=True)
        
        return output_dir / filename
