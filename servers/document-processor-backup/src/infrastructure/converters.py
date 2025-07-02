"""
Concrete implementations of document converters.

These adapters handle the technical details of document conversion
using external libraries.
"""

import logging
import tempfile
import asyncio
import os
from pathlib import Path
from typing import Optional

# External library imports
import markdown
from weasyprint import HTML
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import html2text
from bs4 import BeautifulSoup

from domain.models import (
    ConversionRequest, ConversionResult, Document, DocumentFormat, PDFEngine
)
from domain.ports import DocumentConverter, TemplateRenderer

logger = logging.getLogger(__name__)


class MarkdownConverter(DocumentConverter):
    """Converter for markdown documents."""
    
    def __init__(self, template_renderer: TemplateRenderer):
        self._template_renderer = template_renderer
    
    async def convert(self, request: ConversionRequest) -> ConversionResult:
        """Convert markdown to target format."""
        try:
            if request.target_format == DocumentFormat.HTML:
                return await self._convert_to_html(request)
            elif request.target_format == DocumentFormat.PDF:
                return await self._convert_to_pdf(request)
            elif request.target_format == DocumentFormat.DOCX:
                return await self._convert_to_docx(request)
            elif request.target_format == DocumentFormat.TEXT:
                return await self._convert_to_text(request)
            else:
                return ConversionResult(
                    success=False,
                    error_message=f"Unsupported target format: {request.target_format}"
                )
        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            return ConversionResult(
                success=False,
                error_message=str(e)
            )
    
    def supports_format(self, format: DocumentFormat) -> bool:
        """Check if format is supported."""
        return format in [DocumentFormat.HTML, DocumentFormat.PDF, DocumentFormat.DOCX, DocumentFormat.TEXT]
    
    async def _convert_to_html(self, request: ConversionRequest) -> ConversionResult:
        """Convert markdown to HTML."""
        # Convert markdown to HTML
        html_content = markdown.markdown(
            request.source_content,
            extensions=['extra', 'codehilite', 'toc', 'tables']
        )
        
        # Apply template if specified
        if request.template_name:
            if await self._template_renderer.template_exists(request.template_name):
                html_content = await self._template_renderer.render(
                    request.template_name, html_content
                )
            else:
                logger.warning(f"Template {request.template_name} not found, using basic HTML")
                html_content = self._create_basic_html(html_content)
        else:
            html_content = self._create_basic_html(html_content)
        
        document = Document(
            content=html_content,
            format=DocumentFormat.HTML,
            metadata={"template": request.template_name or "default"}
        )
        
        return ConversionResult(
            success=True,
            document=document,
            engine_used="markdown"
        )
    
    async def _convert_to_pdf(self, request: ConversionRequest) -> ConversionResult:
        """Convert markdown to PDF."""
        engine = request.pdf_engine or PDFEngine.WEASYPRINT
        
        if engine == PDFEngine.WEASYPRINT:
            return await self._convert_to_pdf_weasyprint(request)
        elif engine == PDFEngine.PDFKIT:
            return await self._convert_to_pdf_pdfkit(request)
        elif engine == PDFEngine.REPORTLAB:
            return await self._convert_to_pdf_reportlab(request)
        elif engine == PDFEngine.PANDOC:
            return await self._convert_to_pdf_pandoc(request)
        else:
            return ConversionResult(
                success=False,
                error_message=f"Unsupported PDF engine: {engine}"
            )
    
    async def _convert_to_pdf_weasyprint(self, request: ConversionRequest) -> ConversionResult:
        """Convert to PDF using WeasyPrint."""
        try:
            # First convert to HTML
            html_result = await self._convert_to_html(request)
            if not html_result.success:
                return html_result
            
            # Convert HTML to PDF using WeasyPrint
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                HTML(string=html_result.document.content).write_pdf(temp_file.name)
                
                # Read the PDF content
                with open(temp_file.name, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Clean up temp file
                os.unlink(temp_file.name)
            
            document = Document(
                content=pdf_content.decode('latin-1'),  # Store as string for consistency
                format=DocumentFormat.PDF,
                metadata={"engine": "weasyprint"}
            )
            
            return ConversionResult(
                success=True,
                document=document,
                engine_used="weasyprint"
            )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"WeasyPrint conversion failed: {str(e)}"
            )
    
    async def _convert_to_pdf_pdfkit(self, request: ConversionRequest) -> ConversionResult:
        """Convert to PDF using pdfkit."""
        try:
            import pdfkit
            
            # First convert to HTML
            html_result = await self._convert_to_html(request)
            if not html_result.success:
                return html_result
            
            # Convert HTML to PDF using pdfkit
            pdf_content = pdfkit.from_string(html_result.document.content, False)
            
            document = Document(
                content=pdf_content.decode('latin-1'),
                format=DocumentFormat.PDF,
                metadata={"engine": "pdfkit"}
            )
            
            return ConversionResult(
                success=True,
                document=document,
                engine_used="pdfkit"
            )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"pdfkit conversion failed: {str(e)}"
            )
    
    async def _convert_to_pdf_reportlab(self, request: ConversionRequest) -> ConversionResult:
        """Convert to PDF using ReportLab."""
        try:
            # Convert markdown to plain text for ReportLab
            html_content = markdown.markdown(request.source_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                # Split text into paragraphs
                paragraphs = text_content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        p = Paragraph(para.strip(), styles['Normal'])
                        story.append(p)
                        story.append(Spacer(1, 12))
                
                doc.build(story)
                
                # Read the PDF content
                with open(temp_file.name, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Clean up temp file
                os.unlink(temp_file.name)
            
            document = Document(
                content=pdf_content.decode('latin-1'),
                format=DocumentFormat.PDF,
                metadata={"engine": "reportlab"}
            )
            
            return ConversionResult(
                success=True,
                document=document,
                engine_used="reportlab"
            )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"ReportLab conversion failed: {str(e)}"
            )
    
    async def _convert_to_pdf_pandoc(self, request: ConversionRequest) -> ConversionResult:
        """Convert to PDF using Pandoc."""
        try:
            # Create temporary input file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as temp_input:
                temp_input.write(request.source_content)
                temp_input_path = temp_input.name
            
            # Create temporary output file
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_output:
                temp_output_path = temp_output.name
            
            # Run pandoc
            cmd = ['pandoc', temp_input_path, '-o', temp_output_path, '--pdf-engine=xelatex']
            
            process = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            
            stdout, stderr = await process.communicate()
            
            # Clean up input file
            os.unlink(temp_input_path)
            
            if process.returncode == 0:
                # Read the PDF content
                with open(temp_output_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                
                # Clean up output file
                os.unlink(temp_output_path)
                
                document = Document(
                    content=pdf_content.decode('latin-1'),
                    format=DocumentFormat.PDF,
                    metadata={"engine": "pandoc"}
                )
                
                return ConversionResult(
                    success=True,
                    document=document,
                    engine_used="pandoc"
                )
            else:
                # Clean up output file
                if os.path.exists(temp_output_path):
                    os.unlink(temp_output_path)
                
                return ConversionResult(
                    success=False,
                    error_message=f"Pandoc failed: {stderr.decode()}"
                )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"Pandoc conversion failed: {str(e)}"
            )
    
    async def _convert_to_docx(self, request: ConversionRequest) -> ConversionResult:
        """Convert markdown to DOCX."""
        try:
            # Convert markdown to HTML first
            html_content = markdown.markdown(request.source_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            doc = DocxDocument()
            
            # Process HTML elements
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol']):
                if element.name.startswith('h'):
                    # Headers
                    level = int(element.name[1])
                    doc.add_heading(element.get_text(), level=level)
                elif element.name == 'p':
                    # Paragraphs
                    doc.add_paragraph(element.get_text())
                elif element.name in ['ul', 'ol']:
                    # Lists
                    for li in element.find_all('li'):
                        style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        doc.add_paragraph(li.get_text(), style=style)
            
            # Save to temporary file and read content
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                doc.save(temp_file.name)
                
                with open(temp_file.name, 'rb') as docx_file:
                    docx_content = docx_file.read()
                
                # Clean up temp file
                os.unlink(temp_file.name)
            
            document = Document(
                content=docx_content.decode('latin-1'),
                format=DocumentFormat.DOCX,
                metadata={"engine": "python-docx"}
            )
            
            return ConversionResult(
                success=True,
                document=document,
                engine_used="python-docx"
            )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"DOCX conversion failed: {str(e)}"
            )
    
    async def _convert_to_text(self, request: ConversionRequest) -> ConversionResult:
        """Convert markdown to plain text."""
        try:
            # Convert to HTML first, then to text
            html_content = markdown.markdown(request.source_content)
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = False
            text_content = h.handle(html_content)
            
            document = Document(
                content=text_content,
                format=DocumentFormat.TEXT,
                metadata={"engine": "html2text"}
            )
            
            return ConversionResult(
                success=True,
                document=document,
                engine_used="html2text"
            )
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"Text conversion failed: {str(e)}"
            )
    
    def _create_basic_html(self, html_content: str) -> str:
        """Create basic HTML template."""
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #ddd; margin: 0; padding-left: 20px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
